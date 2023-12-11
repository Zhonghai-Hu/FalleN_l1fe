import requests
import json
import docx
import time

def translate(word):

    """docxFile = docx.Document(path)
    paragraphs = docxFile.paragraphs
    docxNew = docx.Document()"""

    """for par in paragraphs:
        par_text = par.text
        if par_text != "":
            word = {"i":par_text, "doctype":"json"}
            time.sleep(1)
            response = requests.post(url , data = word , headers = header)
            data = json.loads(response.text)
            translate = data["translateResult"][0][0]["tgt"]
            print(translate)
            docxNew.add_paragraph(translate)"""

    url = "https://dict.youdao.com/webtranslate"

    header = {"User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/118.0.0.0 Safari/537.36"}

    word = {"i":word, "doctype":"json"}
    time.sleep(1)
    response = requests.post(url , data = word , headers = header)
    data = json.loads(response.text)
    translate = data["msg"]
    print(translate)

    #docxNew.save(savePath)
translate("hello world")
